from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = ROOT / "outputs" / "figures"
OUTPUT = REPORTS / "이변_예측_최종_분석_보고서.docx"

# standard_business_brief preset with a named Korean typography override.
FONT = "Malgun Gothic"
NAVY = "183B56"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "69737D"
GOLD = "B68118"
PALE_GOLD = "FFF6DC"
RED = "9B1C1C"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
BLACK = "1F2328"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Table widths must total {TABLE_WIDTH}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, color=BLACK, italic=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bottom_border(paragraph, color=BLUE, size=10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_field(paragraph, instruction: str, display: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AI 동아리 토이 프로젝트  |  이변 분석")
    set_run_font(r, size=8.5, bold=True, color=MID_GRAY)
    add_bottom_border(p, color="D7DEE5", size=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    add_field(p, "PAGE", "1")


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, label_color=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    a = p.add_run(f"{label}  ")
    set_run_font(a, size=11, bold=True, color=label_color)
    b = p.add_run(text)
    set_run_font(b, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size=8.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(value)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 1, 2) else WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(str(value))
            set_run_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, filename: str, caption: str, width=6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MID_GRAY)


def build() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_header_footer(section)

    # Editorial cover pattern.
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("FINAL ANALYSIS REPORT")
    set_run_font(kr, size=10, bold=True, color=GOLD)
    kicker.paragraph_format.space_after = Pt(16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("이변 예측 모델 및\n수익률 검증 결과")
    set_run_font(tr, size=30, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(10)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("비인기마 입상 가능성 선별과 연승 베팅 ROI 분석")
    set_run_font(sr, size=14, color=DARK_BLUE)
    sub.paragraph_format.space_after = Pt(38)
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(rule, color=GOLD, size=10)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("AI 동아리 토이 프로젝트 · 2026년 8월 22일")
    set_run_font(mr, size=11, bold=True, color=NAVY)
    meta.paragraph_format.space_before = Pt(16)
    meta.paragraph_format.space_after = Pt(8)
    lock = doc.add_paragraph()
    lock.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lock.add_run("잠금 설정 후 test 1회 평가 · test 확인 후 재튜닝 금지")
    set_run_font(lr, size=9.5, italic=True, color=MID_GRAY)
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    add_callout(
        doc,
        "핵심 결론",
        "Random Forest는 비인기마 중 입상 가능성이 높은 후보를 유의미하게 압축했다. 그러나 실현 ROI는 고배당 한 건에 크게 의존하므로 현재 결과만으로 양의 수익성을 확정할 수 없다.",
    )
    add_bullet(doc, "주 모델: 다크호스 Core Random Forest 600 trees (max_depth=8, min_samples_leaf=30, max_features=sqrt)")
    add_bullet(doc, "잠금 test 성능: ROC-AUC 0.647, PR-AUC 0.213")
    add_bullet(doc, "예측 점수 상위 10% 입상률: 24.7% (전체 13.5% 대비 Lift 1.82)")
    add_bullet(doc, "상위 10% 예상 ROI +39.0%, 실현 ROI +71.4%")
    add_bullet(doc, "최고 수익 한 건 제거 시 상위 10% 실현 ROI -2.3%; 95% 신뢰구간도 0을 포함")
    add_callout(
        doc,
        "판단",
        "모델의 이변 후보 선별력은 재현됐지만, 수익성은 아직 통계적으로 확정되지 않았다. 현재 모델은 베팅 자동화보다 후보 랭킹 도구로 사용하는 것이 타당하다.",
        fill=PALE_GOLD,
        label_color=GOLD,
    )

    doc.add_heading("2. 분석 설계", level=1)
    doc.add_heading("2.1 타깃 정의", level=2)
    add_body(doc, "다크호스: 인기 하위 50%(pop_pct >= 0.50) 중 입상(place == 1)한 말")
    add_body(doc, "인기마 붕괴: 인기 상위 25%(pop_pct <= 0.25) 중 착순 하위 50%(fin_pct >= 0.50)인 말")
    doc.add_heading("2.2 데이터와 검증 원칙", level=2)
    add_bullet(doc, "시간 순서에 따라 train / valid / test로 분리")
    add_bullet(doc, "test 6,660행: 다크호스 후보 3,528행(양성 478건), 인기마 후보 1,902행(양성 484건)")
    add_bullet(doc, "저장 라벨과 재계산 라벨 일치율: 두 타깃 모두 100%")
    add_bullet(doc, "다크호스 연승배당: 결측·비양수·999 이상 0건, 중앙값 6.1, 최댓값 295.7")
    add_bullet(doc, "설정 잠금 후 test를 한 번만 평가하고 test 결과에 따른 재튜닝을 금지")
    doc.add_heading("2.3 피처셋", level=2)
    add_body(doc, "Core는 당일 시장·결과·식별자와 과거 시장 파생변수를 제외한다. History+는 당일 시장·결과·식별자는 제외하되 과거 시장 파생변수를 포함한다.")

    doc.add_heading("3. 모델 선정", level=1)
    add_body(doc, "선정 기준은 valid Lift@10%이며, 동률이면 PR-AUC를 사용했다. Random Forest는 24개 조합을 200개 트리로 선별한 뒤 선택 조합을 600개 트리로 재학습하고 5개 시드로 안정성을 확인했다.")
    add_table(
        doc,
        ["타깃", "피처셋", "선정 모델", "ROC-AUC", "PR-AUC", "Lift@10", "5시드 Lift@10"],
        [
            ["darkhorse", "core", "rf_d8_leaf30_mfsqrt", "0.648", "0.215", "1.954", "1.767 ± 0.067"],
            ["darkhorse", "history_plus", "rf_d6_leaf100_mf0.5", "0.656", "0.219", "1.933", "1.938 ± 0.084"],
            ["favorite_bust", "core", "rf_d8_leaf100_mf0.5", "0.611", "0.346", "1.530", "1.530 ± 0.014"],
            ["favorite_bust", "history_plus", "logit_c0.1", "0.600", "0.339", "1.651", "1.651 ± 0.000"],
        ],
        [1450, 1050, 2100, 950, 950, 900, 1960],
        font_size=8,
    )
    add_body(doc, "다크호스 주 모델은 Core Random Forest다. History+는 비교용 보조 모델이다. 인기마 붕괴에서는 Core가 Random Forest, History+는 Logistic Regression이 선택됐다.")

    doc.add_page_break()
    doc.add_heading("4. 잠금 test 성능", level=1)
    add_table(
        doc,
        ["타깃", "피처셋", "모델", "ROC-AUC", "PR-AUC", "기준률", "상위10% 적중률", "Lift@10"],
        [
            ["darkhorse", "core", "rf_d8_leaf30_mfsqrt", "0.647", "0.213", "13.5%", "24.7%", "1.82"],
            ["darkhorse", "history_plus", "rf_d6_leaf100_mf0.5", "0.649", "0.212", "13.5%", "23.3%", "1.72"],
            ["favorite_bust", "core", "rf_d8_leaf100_mf0.5", "0.574", "0.319", "25.4%", "34.7%", "1.37"],
            ["favorite_bust", "history_plus", "logit_c0.1", "0.582", "0.318", "25.4%", "35.8%", "1.41"],
        ],
        [1400, 900, 1900, 850, 850, 800, 1450, 1210],
        font_size=7.8,
    )
    add_figure(doc, "test_lift_by_percentile.png", "그림 1. 잠금 test의 예측 점수 상위 퍼센트별 Lift")
    add_body(doc, "다크호스 Core는 상위 1%에서 Lift 2.95, 상위 10%에서 1.82를 기록했다. 상위 구간이 넓어질수록 Lift가 점진적으로 1에 수렴해 모델 점수가 후보 우선순위로 기능함을 보여준다.")

    doc.add_heading("4.1 주요 모델 변수", level=2)
    add_figure(doc, "darkhorse_core_feature_importance.png", "그림 2. 다크호스 Core Random Forest의 상위 15개 변수")
    add_body(doc, "중요도 상위에는 최근 훈련량, 직전 경주 순위, 레이팅, 마필·조교사·기수의 입상 및 승률, 휴식일과 부담중량 관련 변수가 포함됐다. 이는 시장 인기 변수를 직접 쓰지 않아도 컨디션·기초 능력·관계자 성과 신호가 이변 후보를 구분하는 데 기여했음을 시사한다.")
    add_callout(doc, "주의", "Random Forest의 불순도 기반 중요도는 인과관계를 의미하지 않으며, 연속형·고유값이 많은 변수에 유리할 수 있다.", fill=LIGHT_GRAY, label_color=MID_GRAY)

    doc.add_page_break()
    doc.add_heading("5. 예측 상위 퍼센트별 ROI", level=1)
    add_body(doc, "1단위 연승 베팅을 가정했다. 예상 ROI는 valid에서 Platt 보정한 확률과 최종 연승배당으로 p × 배당 - 1, 실현 ROI는 입상 × 배당 - 1로 계산했다. 신뢰구간은 경주를 군집 단위로 5,000회 부트스트랩했다.")
    add_table(
        doc,
        ["누적 구간", "베팅 수", "적중률", "Lift", "예상 ROI", "실현 ROI", "실현 ROI 95% CI", "최고수익 1건 제거"],
        [
            ["상위 1%", "35", "40.0%", "2.95", "20.2%", "11.4%", "-33.4% ~ 59.4%", "2.6%"],
            ["상위 2%", "70", "31.4%", "2.32", "20.7%", "9.4%", "-30.7% ~ 53.7%", "-1.7%"],
            ["상위 5%", "176", "25.6%", "1.89", "65.9%", "143.5%", "-22.9% ~ 463.9%", "-3.8%"],
            ["상위 10%", "352", "24.7%", "1.82", "39.0%", "71.4%", "-17.6% ~ 237.1%", "-2.3%"],
            ["상위 20%", "705", "22.8%", "1.69", "23.6%", "29.3%", "-18.1% ~ 111.3%", "-7.5%"],
            ["상위 30%", "1,058", "20.0%", "1.48", "19.7%", "8.8%", "-23.8% ~ 65.7%", "-15.8%"],
            ["상위 50%", "1,764", "18.3%", "1.35", "16.5%", "10.1%", "-24.8% ~ 76.2%", "-4.7%"],
            ["상위 100%", "3,528", "13.5%", "1.00", "4.4%", "-8.2%", "-35.3% ~ 41.5%", "-16.6%"],
        ],
        [1050, 730, 850, 650, 930, 930, 2050, 2170],
        font_size=7.5,
    )
    add_figure(doc, "darkhorse_core_roi_by_percentile.png", "그림 3. 다크호스 Core의 예상·실현 ROI와 고배당 민감도")
    add_callout(
        doc,
        "수익률 해석",
        "상위 2~5% 독립 구간에 배당 295.7의 적중 한 건이 포함돼 누적 5~50% 실현 ROI를 크게 끌어올렸다. 이 한 건을 제거하면 대부분 구간이 음수이고 모든 주요 누적 구간의 95% CI가 0을 포함한다.",
        fill=PALE_RED,
        label_color=RED,
    )

    doc.add_heading("6. 고정 운영 규칙", level=1)
    add_body(doc, "valid 상위 10% 점수 임계값을 test에 그대로 적용한 뒤 같은 경주에서는 최고 점수 말 1두만 선택했다.")
    add_table(
        doc,
        ["피처셋", "베팅/경주", "적중", "적중률", "예상 ROI", "실현 ROI", "95% CI", "최고수익 1건 제거"],
        [
            ["core", "323", "79", "24.5%", "38.7%", "76.0%", "-20.1% ~ 247.0%", "-4.3%"],
            ["history_plus", "304", "72", "23.7%", "40.4%", "70.7%", "-28.8% ~ 254.0%", "-14.7%"],
        ],
        [1250, 1050, 750, 950, 1050, 1050, 1750, 1510],
        font_size=8,
    )
    add_callout(doc, "고정 후보 규칙", "Core 점수 >= 0.192489인 후보 중 경주별 최고점 1두를 선택한다.")
    add_body(doc, "현재 배당은 최종 확정 배당이다. 예상 ROI를 실제 의사결정에 사용하려면 베팅 시점 배당 스냅샷이 필요하다.")

    doc.add_heading("7. 최종 해석 및 권고", level=1)
    add_number(doc, "모델은 비인기마 중 입상 가능성이 높은 말을 유의미하게 압축한다. 상위 1~10% Lift는 1.82~2.95다.")
    add_number(doc, "수익률은 고배당 한 건의 영향이 매우 크므로, 이번 결과는 수익성 확정이 아니라 랭킹 모델의 선별력 검증 완료로 해석한다.")
    add_number(doc, "다음 검증에서는 더 긴 기간의 walk-forward 검증, 베팅 시점 배당 저장, 확률 보정 재검증, 배당 구간별 표본 확대가 필요하다.")
    add_number(doc, "잠금 test 결과를 확인한 뒤에는 하이퍼파라미터나 임계값을 바꾸지 않는다. 변경 아이디어는 새 기간 데이터에서 별도 검증한다.")

    doc.add_page_break()
    doc.add_heading("8. 재현성과 산출물", level=1)
    add_bullet(doc, "잠금 설정: configs/locked_config.json")
    add_bullet(doc, "test 평가 표식: configs/test_evaluation_marker.json")
    add_bullet(doc, "학습 모델: outputs/models/")
    add_bullet(doc, "test 예측: outputs/predictions/")
    add_bullet(doc, "모델·ROI·감사 표: outputs/tables/")
    add_bullet(doc, "그래프: outputs/figures/")
    add_bullet(doc, "실행 코드 및 테스트: src/, scripts/, tests/")
    add_body(doc, "잠금 설정 SHA-256: 7076b51812e808082e024ee79b31d1176e11a56433d409472c196b0981bbe1ad")
    add_callout(doc, "최종 상태", "설정 잠금 후 test 1회 평가를 완료했으며, 단위 테스트 6개가 모두 통과했다.", fill=LIGHT_GRAY, label_color=DARK_BLUE)

    # Word requires a trailing paragraph after a final table. Keep it at a
    # minimal height so it cannot create an extra blank page.
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        last = doc.paragraphs[-1]
        last.paragraph_format.space_before = Pt(0)
        last.paragraph_format.space_after = Pt(0)
        last.paragraph_format.line_spacing = Pt(1)
        tiny = last.add_run(" ")
        set_run_font(tiny, size=1, color=WHITE)
    doc.core_properties.title = "이변 예측 모델 및 수익률 검증 결과"
    doc.core_properties.subject = "AI 동아리 토이 프로젝트 이변 분석 최종 보고서"
    doc.core_properties.author = "AI 동아리 토이 프로젝트"
    doc.core_properties.keywords = "이변 예측, Random Forest, Lift, ROI, 경마"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
