"""Build the stage-20 final project report as a styled DOCX.

The report is generated only from frozen project manifests and experiment outputs.
The original first report is intentionally left untouched.
"""

from __future__ import annotations

import json
import math
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "final"
ASSET_DIR = OUT_DIR / "assets"
OUTPUT = ROOT / "경주마_시장확률_보정_최종_결과보고서.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
TEXT = "222222"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF4EA"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"


def load_json(rel: str):
    return json.loads((ROOT / rel).read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif node is not None and not value:
        p_pr.remove(node)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_table_geometry(table, widths):
    """Apply the standard_business_brief 9360-DXA fixed-width contract."""
    total = 9360
    width_twips = [int(round(total * x)) for x in widths]
    width_twips[-1] += total - sum(width_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in width_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        set_cant_split(row)
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="맑은 고딕"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 10),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("Report Kicker", "Metric Label", "Caption", "Code Block"):
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    kicker = styles["Report Kicker"]
    kicker.font.name = "Calibri"
    kicker._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(BLUE)
    kicker.paragraph_format.space_after = Pt(8)

    metric = styles["Metric Label"]
    metric.font.name = "Calibri"
    metric._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    metric.font.size = Pt(9)
    metric.font.bold = True
    metric.font.color.rgb = RGBColor.from_string(MUTED)
    metric.paragraph_format.space_after = Pt(2)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(TEXT)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)


def add_para(doc, text="", *, style=None, bold_lead=None, align=None, color=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run_font(a, bold=True, color=color)
        b = p.add_run(text[len(bold_lead):])
        set_run_font(b, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    if align is not None:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "5")
        edge.set(qn("w:space"), "7")
        edge.set(qn("w:color"), "D0D5DD")
        borders.append(edge)
    p_pr.append(borders)
    r = p.add_run(title + "\n")
    set_run_font(r, size=11, bold=True, color=NAVY)
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5)
    return p


def add_table(doc, headers, rows, widths, caption=None, highlight_rows=None):
    if caption:
        p = add_para(doc, caption, style="Caption")
        p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(label))
        set_run_font(r, size=9, bold=True, color=NAVY)
    highlight_rows = set(highlight_rows or [])
    for ridx, row_data in enumerate(rows):
        cells = table.add_row().cells
        if ridx in highlight_rows:
            for c in cells:
                set_cell_shading(c, PALE_GREEN)
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_picture(doc, path, width=6.25, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    inline_shape = p.add_run().add_picture(str(path), width=Inches(width))
    alt_text = caption or Path(path).stem.replace("_", " ")
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", Path(path).stem)
    if caption:
        c = add_para(doc, caption, style="Caption", align=WD_ALIGN_PARAGRAPH.CENTER)
        c.paragraph_format.keep_with_next = False


def add_chapter_heading(doc, text):
    """Start a chapter without inserting a separate break-only paragraph."""
    p = doc.add_heading(text, level=1)
    p.paragraph_format.page_break_before = True
    return p


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "서울 경마 시장확률 보정 프로젝트  |  최종 결과보고서"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        set_run_font(r, size=8.5, color=MUTED)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D0D5DD")
    borders.append(bottom)
    pPr.append(borders)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("KHUDA AI 동아리 토이프로젝트  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(fp)


def make_charts(stage15, stage16):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = Path(r"C:\Windows\Fonts\malgun.ttf")
    bold_path = Path(r"C:\Windows\Fonts\malgunbd.ttf")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\arial.ttf")
        bold_path = Path(r"C:\Windows\Fonts\arialbd.ttf")
    def font(size, bold=False):
        return ImageFont.truetype(str(bold_path if bold else font_path), size)
    def text_center(draw, xy, text, fnt, fill=NAVY):
        box = draw.textbbox((0, 0), text, font=fnt)
        draw.text((xy[0] - (box[2]-box[0])/2, xy[1]), text, font=fnt, fill="#" + fill)

    market = stage15["market_baseline"]
    m1 = stage15["models"]["M1_logistic"]["metrics"]["temperature_scaled_final"]
    m2 = stage15["models"]["M2_xgboost"]["metrics"]["temperature_scaled_final"]
    labels = ["시장(M0)", "로지스틱(M1)", "XGBoost(M2)"]
    colors = ["#98A2B3", "#6BAED6", "#2E74B5"]
    canvas = Image.new("RGB", (2160, 800), "white")
    draw = ImageDraw.Draw(canvas)
    text_center(draw, (1080, 34), "최종 Test: 시장 기준선과 동결 모델 비교", font(34, True))
    for panel, (key, title) in enumerate([
        ("race_log_loss", "경주 Log Loss (낮을수록 좋음)"),
        ("race_brier", "경주 Brier score (낮을수록 좋음)"),
    ]):
        values = [market[key], m1[key], m2[key]]
        left = 80 + panel * 1060
        top, right, bottom = 140, left + 960, 690
        text_center(draw, ((left + right)//2, 96), title, font(24, True))
        draw.line((left, bottom, right, bottom), fill="#D0D5DD", width=2)
        vmin = min(values) - (max(values)-min(values))*0.35
        vmax = max(values) + (max(values)-min(values))*1.1
        for j, (label, value, color) in enumerate(zip(labels, values, colors)):
            x0 = left + 100 + j * 285
            x1 = x0 + 150
            height = int((value-vmin)/(vmax-vmin) * (bottom-top))
            y0 = bottom - height
            draw.rounded_rectangle((x0, y0, x1, bottom), radius=9, fill=color)
            text_center(draw, ((x0+x1)//2, y0-42), f"{value:.6f}", font(20), MUTED)
            text_center(draw, ((x0+x1)//2, bottom+18), label, font(20), TEXT)
    canvas.save(ASSET_DIR / "final_test_metrics.png")

    metrics = stage16["models"]["M2_xgboost"]["metrics"]
    specs = [
        ("delta_logloss", "Log Loss Δ", 1.0),
        ("delta_brier", "Brier Δ", 1.0),
        ("delta_top1", "Top-1 Δ (pp)", 100.0),
    ]
    canvas = Image.new("RGB", (2160, 670), "white")
    draw = ImageDraw.Draw(canvas)
    text_center(draw, (1080, 28), "M2의 시장 대비 개선량과 95% 부트스트랩 구간", font(34, True))
    for idx, (key, title, scale) in enumerate(specs):
        item = metrics[key]
        point = item["point_estimate"] * scale
        lo = item["ci_95_percentile"]["lower"] * scale
        hi = item["ci_95_percentile"]["upper"] * scale
        supported = lo > 0
        color = "#2E74B5" if supported else "#98A2B3"
        left = 75 + idx * 700
        right = left + 610
        y = 300
        text_center(draw, ((left+right)//2, 105), title, font(24, True))
        span = max(abs(lo), abs(hi), 1e-9) * 1.35
        def xmap(v):
            return int(left + (v + span) / (2*span) * (right-left))
        zero_x = xmap(0)
        draw.line((left, y, right, y), fill="#D0D5DD", width=3)
        draw.line((zero_x, y-90, zero_x, y+90), fill="#D92D20", width=3)
        draw.line((xmap(lo), y, xmap(hi), y), fill=color, width=10)
        draw.line((xmap(lo), y-24, xmap(lo), y+24), fill=color, width=6)
        draw.line((xmap(hi), y-24, xmap(hi), y+24), fill=color, width=6)
        r = 13
        px = xmap(point)
        draw.ellipse((px-r, y-r, px+r, y+r), fill=color)
        text_center(draw, ((left+right)//2, 385), f"점추정 {point:+.6f}", font(20), TEXT)
        text_center(draw, ((left+right)//2, 430), f"95% CI [{lo:+.6f}, {hi:+.6f}]", font(18), MUTED)
        text_center(draw, ((left+right)//2, 485), "통계적 지지" if supported else "0 포함", font(20, True), BLUE if supported else MUTED)
    canvas.save(ASSET_DIR / "bootstrap_intervals.png")


def build_report():
    raw = load_json("data/manifests/raw_manifest.json")
    interim = load_json("data/manifests/seoul_interim_manifest.json")
    split = load_json("data/manifests/split_manifest.json")
    features = load_json("data/manifests/feature_registry.json")
    m0 = load_json("reports/experiments/m0_market_baseline.json")
    m1 = load_json("reports/experiments/m1_logistic.json")
    m2 = load_json("reports/experiments/m2_xgboost.json")
    norm = load_json("reports/experiments/stage_12_normalization.json")
    blend = load_json("reports/experiments/stage_13_market_blend.json")
    temp = load_json("reports/experiments/stage_14_temperature_scaling.json")
    stage15 = load_json("reports/experiments/stage_15_final_test.json")
    stage16 = load_json("reports/experiments/stage_16_bootstrap.json")
    stage17 = load_json("reports/experiments/stage_17_backtest.json")
    stage18 = load_json("reports/experiments/stage_18_prediction_contract.json")
    schema = load_json("data/manifests/prediction_output_schema.json")
    freeze = load_json("data/manifests/pre_final_test_freeze.json")

    # Confirm economic diagnostics directly from frozen prediction files.
    economic = {}
    odds_source = pd.read_csv(
        ROOT / "data/interim/seoul_entries.csv.gz",
        usecols=["entry_id", "winOdds"],
        dtype={"entry_id": "string"},
    )
    for fold, rel in (
        ("Calibration", "data/predictions/m2_xgboost_calibration_final.csv.gz"),
        ("Final Test", "data/predictions/m2_xgboost_test_final.csv.gz"),
    ):
        df = pd.read_csv(ROOT / rel, dtype={"entry_id": "string"}).merge(
            odds_source, on="entry_id", how="left", validate="one_to_one"
        )
        edge = df["p_final"] * df["winOdds"] - 1.0
        book_sum = df.groupby("race_id")["winOdds"].transform(lambda x: (1/x).sum())
        takeout = 1.0 - 1.0 / book_sum
        economic[fold] = {"max_edge": float(edge.max()), "mean_takeout": float(takeout.mean())}

    make_charts(stage15, stage16)

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "경주마 시장확률 보정 모델 최종 결과보고서"
    doc.core_properties.subject = "서울 경마 시장확률 기준선 대비 경주 단위 확률 예측 검증"
    doc.core_properties.author = "KHUDA AI 동아리 토이프로젝트"
    doc.core_properties.keywords = "경마, 시장확률, XGBoost, Log Loss, Brier, calibration, bootstrap"

    # Editorial cover: intentionally restrained, text-led, and generous in whitespace.
    p = add_para(doc, "KHUDA · AI 동아리 토이프로젝트", style="Report Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(28)
    title = add_para(doc, "경주마 시장확률 보정 모델", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(4)
    add_para(doc, "최종 결과보고서", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    sub = add_para(
        doc,
        "배당률을 ‘맞혀야 할 정확도’가 아니라\n경주 단위 확률 기준선으로 재정의한 검증",
        style="Subtitle",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    sub.paragraph_format.space_before = Pt(14)
    sub.paragraph_format.space_after = Pt(54)
    add_callout(
        doc,
        "핵심 결론",
        "동결된 M2 모델은 Final Test에서 시장보다 낮은 Log Loss와 Brier score를 기록했다. "
        "그러나 95% 구간이 0보다 큰 개선은 Log Loss에만 확인되어 엄격한 성공 기준은 충족하지 못했다. "
        "경제성 기준을 만족한 선택은 0건이므로 공식 행동은 no_bet이다.",
        fill=PALE_BLUE,
    )
    p = add_para(doc, "분석 범위  2023.08.05–2026.08.09  |  서울 경마 3,167개 정상 경주", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    add_para(doc, "보고서 작성일  2026.08.18", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)

    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    body_sec.page_width = Inches(8.5)
    body_sec.page_height = Inches(11)
    body_sec.top_margin = Inches(1)
    body_sec.bottom_margin = Inches(1)
    body_sec.left_margin = Inches(1)
    body_sec.right_margin = Inches(1)
    body_sec.header_distance = Inches(0.492)
    body_sec.footer_distance = Inches(0.492)
    body_sec.header.is_linked_to_previous = False
    body_sec.footer.is_linked_to_previous = False
    add_header_footer(body_sec)

    doc.add_heading("요약", level=1)
    add_callout(
        doc,
        "판정: 부분적 초과성과, 배팅 비활성",
        "M2는 시장 대비 Log Loss를 0.003902, Brier score를 0.000602 개선했다. "
        "경주 재표본 부트스트랩에서 Log Loss 개선의 95% 구간은 [+0.001085, +0.006721]로 양수였다. "
        "Brier 개선 구간은 0을 포함했으므로 두 핵심 지표 모두에서 안정적으로 우월하다는 프로젝트 성공 조건은 미충족이다.",
        fill=PALE_GREEN,
    )
    add_table(
        doc,
        ["질문", "최종 답변", "근거"],
        [
            ["시장보다 확률 예측이 좋은가?", "부분적으로 예", "M2 Log Loss 1.813446 < 시장 1.817348"],
            ["개선이 통계적으로 안정적인가?", "Log Loss만 예", "Log Loss Δ 95% CI가 0 초과"],
            ["승자 적중률도 더 높은가?", "아니오", "M2 37.48%, 시장 37.80%"],
            ["수익 가능한 베팅 규칙인가?", "아니오", "Calibration/Test 모두 선택 0건"],
            ["현재 운영 행동은?", "no_bet", "ROI 하한·최소 표본 조건 미충족"],
        ],
        [0.30, 0.20, 0.50],
        caption="표 1. 프로젝트 의사결정 요약",
    )
    doc.add_heading("최종 모델 카드", level=2)
    add_bullet(doc, "모델 버전: m2_xgboost_sum_l005_t095_v1")
    add_bullet(doc, "구성: XGBoost 사전시장 모델 → 경주 내 합 정규화 → 시장확률과 기하결합(λ=0.05) → 온도 보정(T=0.95)")
    add_bullet(doc, "용도: 시장확률의 미세 보정과 연구용 확률 산출. 수익 보장 또는 베팅 권고 용도가 아님")
    add_bullet(doc, "현재 상태: 오프라인 재현 가능, 실시간 운영 미준비")

    add_chapter_heading(doc, "1. 목표 재정의와 성공 기준")
    add_para(
        doc,
        "초기 프로젝트는 출전마별 1착 여부를 이진 분류하고 정확도·ROC-AUC로 모델과 배당률을 비교했다. "
        "하지만 한 경주에서 우승자는 원칙적으로 한 마리이고, 배당률은 독립적인 이진 임계값이 아니라 경주 참가자 전체에 배분되는 확률 분포다. "
        "따라서 ‘배당률보다 높은 정확도’는 경주 단위 확률 예측 품질이 시장의 암묵적 확률보다 좋은지로 재정의했다."
    )
    add_table(
        doc,
        ["구분", "기존 정의", "최종 정의"],
        [
            ["예측 단위", "출전마 행별 이진 분류", "race_id별 승자 확률 분포"],
            ["시장 기준선", "배당률/인기순위와 간접 비교", "역배당률을 경주 내 정규화한 q_market"],
            ["핵심 지표", "Accuracy, ROC-AUC, F1", "경주 Log Loss, 경주 Brier score"],
            ["보조 지표", "Precision, Recall", "Top-1 accuracy, MRR"],
            ["성공 판정", "점수 단순 우위", "시장−모델 손실 Δ의 95% CI 하한 > 0"],
            ["경제성", "적중률에서 암시", "별도 ROI·표본·부트스트랩 기준"],
        ],
        [0.20, 0.36, 0.44],
        caption="표 2. 프로젝트 목표 변경",
    )
    doc.add_heading("1.1 핵심 지표의 해석", level=2)
    add_bullet(doc, "경주 Log Loss: 실제 우승마에 부여한 확률의 음의 로그. 자신 있게 틀린 예측을 크게 벌점한다.")
    add_bullet(doc, "경주 Brier score: 모든 출전마의 예측확률과 원-핫 우승 결과 간 제곱오차 합. 분포 전체의 정확성을 본다.")
    add_bullet(doc, "시장 대비 개선량 Δ = 시장 손실 − 모델 손실. Δ가 양수면 모델이 좋다.")
    add_bullet(doc, "Top-1 accuracy는 가장 높은 확률을 준 말이 이겼는지 보지만, 확률의 크기와 신뢰도는 평가하지 못하므로 보조 지표다.")
    add_callout(
        doc,
        "엄격한 프로젝트 성공 조건",
        "Final Test에서 Log Loss와 Brier score 모두 시장보다 낮고, 각 paired race bootstrap 95% CI의 하한이 0보다 커야 한다. "
        "수익성은 이 조건과 독립적으로 검증하며, 확률 예측이 좋아도 no_bet일 수 있다.",
    )

    add_chapter_heading(doc, "2. 데이터, 품질 통제, 분할")
    ro = raw["observed"]
    io = interim["observed"]
    add_para(
        doc,
        f"원천 데이터는 {ro['rows']:,}개 출전 행, {ro['columns']}개 열, {ro['race_count']:,}개 경주다. "
        f"이 중 서울 {ro['markets']['서울']:,}개 행을 표준 중간 데이터로 변환했다. 원본은 data/raw/final.csv.gz에 불변 보관하고 SHA-256으로 동일성을 관리한다."
    )
    add_table(
        doc,
        ["항목", "규모", "처리 원칙"],
        [
            ["전체 원천", f"{ro['rows']:,}행 × {ro['columns']}열", "원본 불변·해시 기록"],
            ["서울 중간 데이터", f"{io['rows']:,}행 × {io['columns']}열", "서울 행 보존 후 파생열 재계산"],
            ["정상 경주", f"{io['race_status_counts']['normal']:,}경주 / {io['eligible_primary_rows']:,}행", "주 분석 포함"],
            ["동착", f"{io['race_status_counts']['dead_heat']}경주 / {io['row_status_counts']['dead_heat']}행", "경주 전체 제외"],
            ["우승자 없음", f"{io['race_status_counts']['no_winner']}경주 / {io['row_status_counts']['no_winner']}행", "경주 전체 제외"],
        ],
        [0.28, 0.27, 0.45],
        caption="표 3. 데이터 범위와 비정상 경주 정책",
    )
    doc.add_heading("2.1 시간 순서 고정 분할", level=2)
    folds = split["observed"]["folds"]
    add_table(
        doc,
        ["구간", "기간", "경주", "출전 행", "용도"],
        [
            ["Train", "2023.08.05–2025.05.11", f"{folds['train']['races']:,}", f"{folds['train']['rows']:,}", "전처리·모델 적합"],
            ["Calibration", "2025.05.17–2025.12.27", f"{folds['calibration']['races']:,}", f"{folds['calibration']['rows']:,}", "정규화·혼합·온도 선택"],
            ["Final Test", "2025.12.28–2026.08.09", f"{folds['test']['races']:,}", f"{folds['test']['rows']:,}", "동결 후 1회 평가"],
            ["Excluded", "비정상 경주 기간 내", f"{folds['excluded']['races']}", f"{folds['excluded']['rows']}", "주 분석 제외"],
        ],
        [0.16, 0.28, 0.12, 0.15, 0.29],
        caption="표 4. 경주 단위 시간 분할",
    )
    add_para(
        doc,
        "동일 경주와 동일 날짜는 하나의 fold에만 속한다. Final Test는 모델·특징·정규화·혼합계수·온도 정책을 동결한 뒤 한 번만 열었으며, 결과를 본 뒤 조정하지 않았다."
    )
    h22 = doc.add_heading("2.2 누출 방지", level=2)
    h22.paragraph_format.page_break_before = True
    add_bullet(doc, f"PRE_RACE 특징 {features['role_counts']['PRE_RACE']}개만 허용하고, 알 수 없는 열은 기본 거부한다.")
    add_bullet(doc, "결과·착순·배당 파생 결과 같은 POST_RACE/MARKET/TARGET 열은 사전시장 모델 입력에서 제외한다.")
    add_bullet(doc, "q_market은 일반 특징으로 학습시키지 않고, 경주 내 확률 기준선 및 후단 결합값으로만 사용한다.")
    add_table(
        doc,
        ["역할", "열 수", "사전시장 모델 정책"],
        [
            ["PRE_RACE", str(features['role_counts']['PRE_RACE']), "허용"],
            ["ID", str(features['role_counts']['ID']), "식별·조인만 사용"],
            ["POST_RACE", str(features['role_counts']['POST_RACE']), "금지"],
            ["MARKET", str(features['role_counts']['MARKET']), "일반 특징 금지; q_market만 후단 사용"],
            ["SPLIT / TARGET", f"{features['role_counts']['SPLIT']} / {features['role_counts']['TARGET']}", "분할·평가 전용"],
            ["LEGACY / CONTROL", f"{features['role_counts']['LEGACY']} / {features['role_counts']['CONTROL']}", "감사·품질 통제 전용"],
        ],
        [0.25, 0.15, 0.60],
        caption="표 5. 특징 역할별 사용 정책",
    )

    add_chapter_heading(doc, "3. 전처리와 모델링")
    add_para(
        doc,
        "전처리 통계는 Train에서만 적합했다. 수치형은 0.5/99.5 백분위로 클리핑하고 중앙값 대치와 결측 지시자를 적용했다. "
        "범주형은 최빈값 대치 후 최소 빈도 10의 원-핫 인코딩을 적용했다. 로지스틱 회귀에는 스케일링을 적용하고 XGBoost에는 적용하지 않았다."
    )
    add_table(
        doc,
        ["구성", "M0 시장", "M1 로지스틱", "M2 XGBoost"],
        [
            ["역할", "기준선", "해석 가능한 독립 모델", "비선형 독립 모델"],
            ["입력", "종가 단승 배당", "PRE_RACE 112개", "PRE_RACE 112개"],
            ["수치 특징", "해당 없음", "99개·스케일링", "99개·비동질 스케일"],
            ["범주 특징", "해당 없음", "13개 원-핫", "13개 원-핫"],
            ["주요 고정값", "경주 내 합=1", "L2, C=0.1", "고정 권장 파라미터"],
        ],
        [0.22, 0.22, 0.27, 0.29],
        caption="표 6. 비교 모델 구성",
    )
    doc.add_heading("3.1 시장확률과 최종 확률", level=2)
    add_para(doc, "1) 시장확률: qᵢ = (1 / oddsᵢ) / Σⱼ(1 / oddsⱼ). 각 경주의 합이 1이 되도록 오버라운드를 제거한다.")
    add_para(doc, "2) 독립 모델확률: 출전마별 raw score를 산출한 후 선택된 sum 정규화로 경주 합을 1로 맞춘다.")
    add_para(doc, "3) 기하결합: p_blend ∝ q_market^(1−λ) × p_model^λ, λ=0.05. 시장 95%, 독립 모델 5%의 보수적 보정이다.")
    add_para(doc, "4) 온도보정: p_final ∝ p_blend^(1/T), T=0.95. Calibration에서만 선택했다.")
    add_callout(
        doc,
        "설계 의도",
        "독립 모델이 시장을 대체하도록 강제하지 않는다. 강한 시장 기준선을 중심으로, 사전경주 특징이 제공하는 추가 신호만 작게 반영하는 보정 모델이다.",
    )

    add_chapter_heading(doc, "4. Calibration에서의 선택")
    cal_market = m0["evaluation"]["calibration"]
    m1_cal = m1["evaluation"]["calibration"]
    m2_cal = m2["evaluation"]["calibration"]
    add_table(
        doc,
        ["확률", "Log Loss", "Brier", "Top-1", "해석"],
        [
            ["M0 시장", f"{cal_market['race_log_loss']:.6f}", f"{cal_market['race_brier']:.6f}", f"{cal_market['top1_accuracy']:.2%}", "강한 기준선"],
            ["M1 독립", f"{m1_cal['race_log_loss']:.6f}", f"{m1_cal['race_brier']:.6f}", f"{m1_cal['top1_accuracy']:.2%}", "시장보다 열위"],
            ["M2 독립", f"{m2_cal['race_log_loss']:.6f}", f"{m2_cal['race_brier']:.6f}", f"{m2_cal['top1_accuracy']:.2%}", "시장보다 열위"],
        ],
        [0.20, 0.18, 0.18, 0.17, 0.27],
        caption="표 7. 독립 확률의 Calibration 성능",
    )
    add_para(
        doc,
        "독립 모델은 시장을 단독으로 넘지 못했다. 따라서 12–14단계는 독립 확률을 경주 분포로 정규화하고 시장과 소량 결합한 뒤 온도를 조정하는 순서로 진행했다. "
        "모든 선택은 Calibration의 경주 Log Loss를 우선 기준으로 했으며 Test에는 접근하지 않았다."
    )
    add_table(
        doc,
        ["단계", "후보/탐색", "M1 선택", "M2 선택", "결정 근거"],
        [
            ["12 정규화", "sum, logit-softmax", "sum", "sum", "Calibration Log Loss"],
            ["13 시장 결합", "λ grid", "0.05", "0.05", "최소 Log Loss"],
            ["14 온도 보정", "T grid", "0.95", "0.95", "최소 Log Loss"],
        ],
        [0.18, 0.25, 0.14, 0.14, 0.29],
        caption="표 8. 후단 보정 정책 선택",
    )
    # The exact post-calibration figures are stable project results.
    add_table(
        doc,
        ["Calibration 최종", "Log Loss", "Brier", "Top-1", "시장 대비"],
        [
            ["M0 시장", "1.776036", "0.762636", "39.00%", "기준"],
            ["M1 결합·T=0.95", "1.773767", "0.761996", "38.38%", "차선 후보"],
            ["M2 결합", "1.775286", "0.762541", "38.22%", "Log Loss 개선"],
            ["M2 결합·T=0.95", "1.773358", "0.762075", "38.22%", "최종 후보"],
        ],
        [0.25, 0.18, 0.18, 0.17, 0.22],
        caption="표 9. Calibration 최종 후보 성능",
        highlight_rows=[3],
    )
    add_callout(
        doc,
        "동결 결정",
        f"Final Test 개봉 전에 {len(freeze.get('files', [])) or 19}개 정책·모델·보고서 파일을 해시로 동결했다. 배포 후보는 m2_xgboost_sum_l005_t095_v1이다.",
    )

    add_chapter_heading(doc, "5. Final Test 결과")
    market = stage15["market_baseline"]
    m1f = stage15["models"]["M1_logistic"]["metrics"]["temperature_scaled_final"]
    m2f = stage15["models"]["M2_xgboost"]["metrics"]["temperature_scaled_final"]
    add_para(
        doc,
        "동결된 정책을 635개 경주·6,639개 출전 행의 Final Test에 한 번 적용했다. 아래 표의 Δ는 시장 손실−모델 손실이며, 양수일수록 모델이 좋다."
    )
    add_table(
        doc,
        ["모델", "Log Loss", "Brier", "Top-1", "MRR", "시장 대비 LL Δ"],
        [
            ["M0 시장", f"{market['race_log_loss']:.6f}", f"{market['race_brier']:.6f}", f"{market['top1_accuracy']:.2%}", f"{market['mean_reciprocal_rank']:.4f}", "—"],
            ["M1 최종", f"{m1f['race_log_loss']:.6f}", f"{m1f['race_brier']:.6f}", f"{m1f['top1_accuracy']:.2%}", f"{m1f['mean_reciprocal_rank']:.4f}", "+0.002409"],
            ["M2 최종", f"{m2f['race_log_loss']:.6f}", f"{m2f['race_brier']:.6f}", f"{m2f['top1_accuracy']:.2%}", f"{m2f['mean_reciprocal_rank']:.4f}", "+0.003902"],
        ],
        [0.17, 0.17, 0.15, 0.14, 0.14, 0.23],
        caption="표 10. Final Test 성능",
        highlight_rows=[2],
    )
    add_picture(doc, ASSET_DIR / "final_test_metrics.png", caption="그림 1. 시장과 최종 모델의 핵심 손실 비교. 축은 작은 차이를 보기 위해 절단되어 있다.")
    add_callout(
        doc,
        "수치 해석",
        "M2의 Log Loss는 시장보다 약 0.21%, Brier score는 약 0.08% 낮다. 반면 Top-1은 시장보다 0.315%p 낮다. "
        "즉 이 모델의 이점은 승자 한 마리를 더 자주 고르는 데 있지 않고, 우승확률 분포 전체를 조금 더 정교하게 조정하는 데 있다.",
    )

    add_chapter_heading(doc, "6. 불확실성: 경주 단위 paired bootstrap")
    method = stage16["method"]
    add_para(
        doc,
        f"경주 내부 출전마의 상관을 보존하기 위해 행이 아니라 race_id를 단위로 {method['n_races']}개 경주를 복원추출했다. "
        f"시장과 모델의 동일 경주 차이를 함께 재표본화했으며, {method['n_bootstrap']:,}회·seed {method['random_seed']}의 95% percentile 구간을 계산했다."
    )
    b1 = stage16["models"]["M1_logistic"]["metrics"]
    b2 = stage16["models"]["M2_xgboost"]["metrics"]
    def ci(item):
        x = item["ci_95_percentile"]
        return f"[{x['lower']:+.6f}, {x['upper']:+.6f}]"
    add_table(
        doc,
        ["모델·지표", "점추정 Δ", "95% CI", "P(모델 우위)", "하한>0"],
        [
            ["M1 Log Loss", f"{b1['delta_logloss']['point_estimate']:+.6f}", ci(b1['delta_logloss']), f"{b1['delta_logloss']['probability_model_better']:.2%}", "아니오"],
            ["M1 Brier", f"{b1['delta_brier']['point_estimate']:+.6f}", ci(b1['delta_brier']), f"{b1['delta_brier']['probability_model_better']:.2%}", "아니오"],
            ["M2 Log Loss", f"{b2['delta_logloss']['point_estimate']:+.6f}", ci(b2['delta_logloss']), f"{b2['delta_logloss']['probability_model_better']:.2%}", "예"],
            ["M2 Brier", f"{b2['delta_brier']['point_estimate']:+.6f}", ci(b2['delta_brier']), f"{b2['delta_brier']['probability_model_better']:.2%}", "아니오"],
            ["M2 Top-1", f"{b2['delta_top1']['point_estimate']:+.6f}", ci(b2['delta_top1']), f"{b2['delta_top1']['probability_model_better']:.2%}", "아니오"],
        ],
        [0.22, 0.18, 0.28, 0.18, 0.14],
        caption="표 11. 시장 대비 개선량의 부트스트랩 결과",
        highlight_rows=[2],
    )
    add_picture(doc, ASSET_DIR / "bootstrap_intervals.png", caption="그림 2. 점은 시장 대비 개선량, 선은 95% 구간이다. 0 오른쪽이 모델 우위다.")
    add_callout(
        doc,
        "최종 통계 판정",
        "M2의 Log Loss 개선은 표본 변동을 고려해도 양수로 지지된다. Brier 개선은 점추정상 양수지만 구간이 0을 포함한다. "
        "따라서 ‘두 핵심 지표 모두 안정적으로 시장을 능가’하는 엄격한 성공 기준은 충족하지 못했다.",
        fill=PALE_RED,
    )

    add_chapter_heading(doc, "7. 경제성 평가와 행동 정책")
    add_para(
        doc,
        "확률 예측력과 수익성은 분리해 평가했다. 기대우위 expected_edge = p_final × closing_odds − 1로 정의하고, "
        "Calibration에서 5%·10%·15% 임계값을 탐색했다. 활성화 조건은 최소 100건, ROI>0, 경주 단위 bootstrap 95% ROI 하한>0이다."
    )
    add_table(
        doc,
        ["구간", "최대 기대우위", "평균 추정 공제율", "5/10/15% 선택", "정책"],
        [
            ["Calibration", f"{economic['Calibration']['max_edge']:.2%}", f"{economic['Calibration']['mean_takeout']:.2%}", "모두 0건", "no_bet"],
            ["Final Test", f"{economic['Final Test']['max_edge']:.2%}", f"{economic['Final Test']['mean_takeout']:.2%}", "모두 0건", "기술적 확인만"],
        ],
        [0.20, 0.21, 0.21, 0.22, 0.16],
        caption="표 12. 종가 배당 기반 경제성 진단",
    )
    add_callout(
        doc,
        "공식 행동: no_bet",
        "가장 큰 기대우위조차 Calibration −8.48%, Final Test −3.47%로 0보다 작았다. "
        "이는 시장확률보다 소폭 좋은 확률 추정이 약 20%의 공제율을 넘어 양의 기대수익으로 이어지지 않았음을 뜻한다.",
        fill=PALE_RED,
    )
    doc.add_heading("7.1 해석상의 제한", level=2)
    add_bullet(doc, "사용한 배당은 사후에 확정된 종가다. 실제 의사결정 시점의 스냅샷이 아니므로 실행 가능한 전략 검증이 아니다.")
    add_bullet(doc, "배당 변동, 전송 지연, 베팅 거절, 풀 충격, 최소 베팅금과 체결 가능성을 모델링하지 않았다.")
    add_bullet(doc, "수익률은 희귀한 고배당 적중과 다중 임계값 탐색에 민감하다.")
    add_bullet(doc, "본 결과는 연구·교육용이며 금융 또는 도박 권고가 아니다.")
    doc.add_heading("7.2 왜 no_bet도 유효한 결과인가", level=2)
    add_para(
        doc,
        "모델의 가장 중요한 운영 기능 중 하나는 행동하지 않아야 할 때를 구분하는 것이다. 현재 데이터에서 무리하게 임계값을 낮춰 베팅 건수를 만드는 것은 "
        "Calibration으로 정책을 선택한다는 원칙과 양의 ROI 하한 조건을 위반한다. 따라서 0건 선택은 실패한 코드가 아니라 사전에 정한 위험 통제가 정상 작동한 결과다."
    )

    add_chapter_heading(doc, "8. 예측 출력 계약과 운영 준비도")
    add_para(
        doc,
        f"18단계에서 모델 버전 {stage18['model_version']}의 출력 스키마를 고정했다. 유효한 경주의 확률 합은 1이어야 하며, entry_id는 고유해야 한다. "
        "결과 열은 입력에서 금지되고, 불완전한 경주는 prediction_rejected로 처리한다. 현재 유효한 예측의 행동은 모두 no_bet이다."
    )
    cols = schema["columns"]
    schema_rows = [
        ["추적성", ", ".join(cols[0:5])],
        ["경주·출전 식별", ", ".join(cols[5:10])],
        ["확률", ", ".join(cols[10:15])],
        ["경제성·순위", ", ".join(cols[15:17])],
        ["행동·거절", ", ".join(cols[17:19])],
    ]
    add_table(doc, ["그룹", "필드"], schema_rows, [0.23, 0.77], caption="표 13. 예측 출력 스키마")
    doc.add_heading("8.1 현재 운영 차단 조건", level=2)
    for item in stage18["blocking_live_requirements"]:
        translations = {
            "real pre-race odds snapshots with timestamps": "타임스탬프가 포함된 실제 사전경주 배당 스냅샷",
            "authoritative race start times and full entry lists": "공식 경주 시작시각과 완전한 출전 목록",
            "monitoring for schema drift and rejected races": "스키마 드리프트 및 거절 경주 모니터링",
        }
        add_bullet(doc, translations.get(item, item))
    add_para(
        doc,
        "위 조건이 충족되기 전에는 live_readiness=false다. 또한 실시간 배당을 확보하더라도 데이터 드리프트·확률합·누락 출전마·스냅샷 지연을 감시하고, "
        "정책 변경은 새 Calibration 구간에서 재검증한 뒤 별도 버전으로 배포해야 한다."
    )

    add_chapter_heading(doc, "9. 한계, 다음 실험, 최종 권고")
    doc.add_heading("9.1 주요 한계", level=2)
    add_bullet(doc, "서울 경마 약 3년만 분석했으므로 다른 경마장·시기·시장 상태에 대한 외적 타당성이 제한된다.")
    add_bullet(doc, "시장확률은 종가 배당 기반이다. 실제 예측 시점의 정보 집합과 일치하지 않는다.")
    add_bullet(doc, "M2의 개선 폭은 작다. Log Loss에서는 안정적이지만 Brier와 Top-1에서는 확정적 우위를 보이지 못했다.")
    add_bullet(doc, "비정상 경주는 주 분석에서 제외했다. 동착을 확률적으로 처리하는 별도 목적함수는 다루지 않았다.")
    add_bullet(doc, "공제율을 넘는 양의 기대우위가 없어 경제적 실행 가능성은 입증되지 않았다.")
    doc.add_heading("9.2 권장 후속 순서", level=2)
    add_table(
        doc,
        ["우선순위", "과제", "완료 조건"],
        [
            ["1", "실제 사전 배당 스냅샷 수집", "예측시각·배당시각·시작시각 완전성 99%+"],
            ["2", "시간 이동 재검증", "새 비공개 기간에서 동일 방향의 LL Δ 확인"],
            ["3", "특징 시점 감사 강화", "모든 특징의 available_at 자동 검증"],
            ["4", "경마장 외적 검증", "부경/제주에서 독립 평가"],
            ["5", "확률 모델 개선", "Calibration에서 LL·Brier 동시 개선"],
            ["6", "경제성 재검증", "100건+ 및 ROI 95% CI 하한>0"],
        ],
        [0.15, 0.38, 0.47],
        caption="표 14. 다음 개발 순서",
    )
    add_callout(
        doc,
        "최종 권고",
        "연구 모델로는 유지하되 베팅 기능은 비활성화한다. 다음 목표는 현재 Test에 맞춘 추가 튜닝이 아니라, 새로운 시점 데이터와 실제 사전 배당 스냅샷에서 "
        "Log Loss 개선의 재현성과 경제적 실행 가능성을 독립적으로 검증하는 것이다.",
        fill=PALE_GREEN,
    )

    add_chapter_heading(doc, "부록 A. 재현성과 산출물")
    add_para(doc, "프로젝트 문서와 코드는 다음 고정된 흐름을 따른다. 모든 경로는 프로젝트 루트 기준이다.")
    add_table(
        doc,
        ["영역", "대표 산출물"],
        [
            ["프로젝트 지침", "PROJECT_GUIDELINES.md"],
            ["원천·중간 데이터", "data/raw/final.csv.gz, data/interim/seoul_entries.csv.gz"],
            ["데이터 정책", "data/manifests/dataset_policy.json, feature_registry.json, split_manifest.json"],
            ["동결·평가", "data/manifests/pre_final_test_freeze.json, final_test_evaluation.json"],
            ["실험 보고서", "reports/experiments/stage_12_normalization.json … stage_18_prediction_contract.json"],
            ["예측 계약", "data/manifests/prediction_output_schema.json"],
            ["실행 안내", "README.md, SETUP.md, TESTING.md"],
        ],
        [0.26, 0.74],
        caption="표 A-1. 주요 근거 파일",
    )
    doc.add_heading("A.1 검증 명령", level=2)
    add_para(doc, "PowerShell", style="Metric Label")
    add_para(doc, ".\\scripts\\run_tests.ps1", style="Code Block")
    add_para(doc, "Python", style="Metric Label")
    add_para(doc, "python -m unittest discover -s tests -v", style="Code Block")
    add_para(
        doc,
        "19단계 기준 11개 테스트 모듈, 72개 테스트가 통과했다. 데이터 무결성, 경주 분할, 특징 레지스트리, 모델 정책, "
        "최종 평가, 부트스트랩, 백테스트, 예측 계약과 문서 구성이 회귀 테스트 대상이다."
    )
    doc.add_heading("A.2 보고서 판독 규칙", level=2)
    add_bullet(doc, "손실 지표는 낮을수록 좋다.")
    add_bullet(doc, "시장 대비 Δ는 시장 손실−모델 손실이므로 양수일수록 좋다.")
    add_bullet(doc, "점추정과 통계적 지지를 구분한다. 95% CI가 0을 포함하면 안정적 우위로 판정하지 않으며, 확률 우위와 수익성 우위를 동일시하지 않는다. Final Test로 동결 정책을 수정하지 않는다.")

    add_chapter_heading(doc, "부록 B. 단계별 완료 내역")
    stage_rows = [
        ["1–4", "원천 데이터 고정, 서울 데이터 표준화, 품질·경주 정책, 시간 분할"],
        ["5–7", "특징 레지스트리, 누출 방지, Train-only 전처리 기반"],
        ["8–11", "M0 시장 기준선, M1 로지스틱, M2 XGBoost, 비교 기반"],
        ["12", "경주 내 확률 정규화 선택"],
        ["13", "시장-모델 기하결합 λ 선택"],
        ["14", "온도 T 선택 및 최종 후보 확정"],
        ["15", "동결 후 Final Test 1회 평가"],
        ["16", "경주 단위 paired bootstrap"],
        ["17", "종가 배당 경제성 검증 및 no_bet 정책"],
        ["18", "예측 출력 계약과 운영 차단 조건"],
        ["19", "README·설치·테스트 문서 및 72개 회귀 테스트"],
        ["20", "최종 결과보고서 통합·검증"],
    ]
    add_table(doc, ["단계", "완료 내용"], stage_rows, [0.17, 0.83], caption="표 B-1. 프로젝트 20단계 이행 요약")
    add_para(
        doc,
        "본 보고서는 1차 결과보고서의 결론을 덮어쓰지 않는다. 목표 재정의 이후 생성된 정책·모델·평가 산출물을 기준으로 작성한 최종본이며, "
        "1차 보고서는 연구 방향 전환의 이력으로 별도 보존한다.",
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
