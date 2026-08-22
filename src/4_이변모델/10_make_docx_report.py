"""
10_make_docx_report.py — 이변 예측 모델 상세 보고서 (docx)

팀원 보고서와 동일한 수준의 상세도로 작성.
08_full_pipeline.py 실행 결과를 읽어서 docx 생성.

실행:
    python src/4_이변모델/10_make_docx_report.py

출력:
    results/upset_with_odds_v2/reports/4_이변모델/04_배당률포함_모델.docx
"""

import logging
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("results/upset_with_odds_v2")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load results
    comp = pd.read_csv(OUTPUT_DIR / "model_comparison.csv") if (OUTPUT_DIR / "model_comparison.csv").exists() else None
    seg = pd.read_csv(OUTPUT_DIR / "segment_comparison.csv") if (OUTPUT_DIR / "segment_comparison.csv").exists() else None
    roi = pd.read_csv(OUTPUT_DIR / "roi_simulation.csv") if (OUTPUT_DIR / "roi_simulation.csv").exists() else None
    ci = pd.read_csv(OUTPUT_DIR / "roi_confidence_interval.csv") if (OUTPUT_DIR / "roi_confidence_interval.csv").exists() else None
    fi = pd.read_csv(OUTPUT_DIR / "feature_importance.csv") if (OUTPUT_DIR / "feature_importance.csv").exists() else None

    doc = Document()

    # ===== Title =====
    title = doc.add_heading("KHUDA 10기 토이프로젝트 3조", level=0)
    doc.add_heading("배당률 포함 이변 예측 모델 보고서", level=1)
    add_para(doc, "배당률을 피처로 활용하면 이변 예측이 얼마나 개선되는가 — A/B/C 비교 실험", italic=True)
    doc.add_paragraph()

    # ===== 0. 요약 =====
    add_heading(doc, "0. 세 줄 요약", level=1)
    add_para(doc, "1. 배당률(q)을 피처에 포함시킨 모델(C)이 배당률 단독(A)보다 이변 예측 성능이 높다 (AUC 0.742 vs 0.737).")
    add_para(doc, "2. C 모델 상위 5% 선별 베팅 시 ROI +25.9% (수수료 반영 후). 부트스트랩 1,000회 중 81%에서 수익 발생.")
    add_para(doc, "3. 다만 95% 신뢰구간이 0을 포함하여, 통계적으로 '확실한 수익'이라 단정하기는 어렵다. 정보 가치는 존재하나 표본 확대가 필요.")
    doc.add_paragraph()

    # ===== 1. 배경 =====
    add_heading(doc, "1. 이 실험의 배경", level=1)
    add_para(doc, "1.1 이전 실험 요약")
    add_para(doc, "우리 프로젝트의 첫 번째 시도는 '배당률을 제외하고 1착을 예측하여 시장을 이기는 것'이었다. 결과는 실패 — 시장(배당률) ROC-AUC 0.81 대 모델 0.75. 어떤 전처리(이상치 제거, 스케일링, 고상관 제거)를 해도 시장을 넘지 못했다.")
    add_para(doc, "그래서 목표를 바꿨다. '시장을 이기겠다'가 아니라 '시장이 놓치는 이변을 배당률과 함께 잡겠다'.")
    doc.add_paragraph()
    add_para(doc, "1.2 이 실험의 핵심 질문")
    add_para(doc, '"배당률(q)을 피처로 포함시키면, 배당률 단독보다 이변을 더 잘 잡을 수 있는가?"', bold=True)
    add_para(doc, "즉, 배당률이 '대략적인 인기도'를 알려주고, 나머지 피처(마필 상태, 기수, 훈련량 등)가 '같은 배당대 안에서 진짜 이길 말'을 세분화해줄 수 있는지를 검증한다.")
    doc.add_paragraph()

    # ===== 2. 실험 설계 =====
    add_heading(doc, "2. 실험 설계", level=1)
    add_para(doc, "2.1 타겟 정의")
    add_para(doc, "upset = 1  if  (pop_pct >= 0.5)  AND  (win == 1)")
    add_para(doc, "pop_pct >= 0.5 = 경주 내 배당률 순위 하위 50% (비인기마). win == 1 = 실제 1착.")
    add_para(doc, "비인기마만을 학습 대상으로 사용. 인기마는 이변 대상이 아니므로 제외.")
    doc.add_paragraph()

    add_para(doc, "2.2 세 가지 모델 비교")
    add_table(doc,
              ["모델", "피처", "의미"],
              [
                  ["A (베이스라인)", "q (배당률 암묵적 확률) 1개", "배당률만으로 이변을 얼마나 잡나?"],
                  ["B (기존)", "말/기수/환경 피처 ~86개 (배당률 완전 제외)", "배당률 없이 이변을 잡을 수 있나?"],
                  ["C (결합)", "q + 말/기수/환경 피처 ~87개", "배당률에 피처를 더하면 더 잘 잡나?"],
              ])
    doc.add_paragraph()

    add_para(doc, "2.3 전처리")
    add_table(doc,
              ["단계", "처리", "비고"],
              [
                  ["1. 필터링", "서울 + 비인기마(pop_pct>=0.5)", "56,648 → 17,392행"],
                  ["2. 결측치", "구조적→0, 랜덤→train 중앙값", "rating, hr_winrate 등 26개 컬럼"],
                  ["3. 다중공선성", "배당률 14개→q만, 고상관 23개 제거", "상관 0.8+ 쌍에서 한쪽 삭제"],
                  ["4. 인코딩", "범주형 13개 → LabelEncoder", "rank, track, weather 등"],
                  ["5. 스케일링", "StandardScaler (train fit)", "로지스틱 회귀를 위해"],
                  ["6. 분할", "시간순 6:2:2", "train 60% / valid 20% / test 20%"],
              ])
    doc.add_paragraph()

    add_para(doc, "2.4 왜 배당률 파생 14개 중 q만 남겼나")
    add_para(doc, "14개 컬럼 간 상관계수를 계산한 결과, q와 나머지 13개의 |r|이 대부분 0.9 이상이었다. 사실상 같은 정보의 다른 표현이므로, 가장 깨끗한 형태인 q(오버라운드 제거 후 정규화, 경주 내 합=1)만 대표로 사용.")
    doc.add_paragraph()

    add_para(doc, "2.5 분할 정보")
    add_table(doc,
              ["Fold", "행 수", "비율", "upset 비율", "기간"],
              [
                  ["train", "10,425", "60%", "2.95%", "2023-08 ~ 2025-05"],
                  ["valid", "3,439", "20%", "2.70%", "2025-05 ~ 2025-12"],
                  ["test", "3,528", "20%", "2.66%", "2025-12 ~ 2026-08"],
              ])
    doc.add_paragraph()

    # ===== 3. 모델 성능 =====
    add_heading(doc, "3. 모델 성능 비교", level=1)

    if comp is not None:
        add_para(doc, "test 3,528행, 기저율 2.66% (94마리가 실제 이변).")
        doc.add_paragraph()

        headers = ["피처셋", "모델", "AUC", "F1-Macro", "Precision", "Recall"]
        rows = []
        for _, r in comp.iterrows():
            rows.append([r["feature_set"], r["model"],
                         f'{r["ROC_AUC"]:.4f}', f'{r["F1_Macro"]:.4f}',
                         f'{r["Precision"]:.4f}', f'{r["Recall"]:.4f}'])
        add_table(doc, headers, rows)
        doc.add_paragraph()

        add_para(doc, "핵심 결과:", bold=True)
        c_auc = comp[comp["feature_set"] == "C (q + features)"]["ROC_AUC"].max()
        a_auc = comp[comp["feature_set"] == "A (q only)"]["ROC_AUC"].max()
        b_auc = comp[comp["feature_set"] == "B (no odds)"]["ROC_AUC"].max()
        add_para(doc, f"  • C(결합) AUC = {c_auc:.4f} > A(배당률만) = {a_auc:.4f} > B(피처만) = {b_auc:.4f}")
        add_para(doc, f"  • C vs A 개선: +{c_auc - a_auc:.4f} — 피처가 배당률 위에 추가 정보를 제공")
        add_para(doc, f"  • A > B — q 하나가 피처 86개보다 강력. 시장 효율성 재확인")
        add_para(doc, "  • RF가 LR보다 우수 — 비선형 관계(훈련량×휴식일 교호작용 등)를 포착")
    doc.add_paragraph()

    # ===== 4. 구간별 =====
    add_heading(doc, "4. 비인기마 배당 구간별 A vs C", level=1)
    add_para(doc, "핵심 질문: '고배당(비인기) 구간에서 C가 A보다 더 잘 잡는가?'")
    add_para(doc, "배당률만으로는 같은 고배당 구간 안에서 말을 구분 못한다. 피처가 추가 정보를 제공하는지 확인.")
    doc.add_paragraph()

    if seg is not None:
        headers = ["구간", "마리 수", "기저율", "A 적중률", "C 적중률", "C-A", "C Lift"]
        rows = []
        for _, r in seg.iterrows():
            rows.append([r["segment"], r["n"], f'{r["baseline"]:.4f}',
                         f'{r["A_hit"]:.4f}', f'{r["C_hit"]:.4f}',
                         f'{r["C_minus_A"]:+.4f}', f'{r["Lift_C"]:.2f}x'])
        add_table(doc, headers, rows)
        doc.add_paragraph()
        add_para(doc, "해석: 10~20배 구간에서 C가 A보다 +0.78%p 우위. 40배+ 구간에서도 +0.33%p. 고배당일수록 '배당률만으로는 부족하고, 피처(훈련량, 기수 컨디션 등)가 추가 정보를 제공'한다는 가설을 지지.")
    doc.add_paragraph()

    # ===== 5. ROI =====
    add_heading(doc, "5. 수익률(ROI) 시뮬레이션", level=1)
    add_para(doc, "5.1 전략 설명")
    add_para(doc, "C 모델이 '이변 확률 높다'고 예측한 상위 N%에만 단승 베팅. winOdds에는 이미 ~20% 공제율이 반영되어 있으므로, 아래 ROI는 수수료 차감 후 실질 수익률이다.")
    doc.add_paragraph()

    add_para(doc, "5.2 세 가지 베팅 전략")
    add_table(doc,
              ["전략", "방식", "장점"],
              [
                  ["Flat", "모든 베팅에 동일 금액", "단순, 비교 기준"],
                  ["Proportional", "확률 높을수록 비례해서 더 많이", "자신 있는 곳에 집중"],
                  ["Kelly (Half)", "기대값 기반 최적 비율 계산", "수학적 최적, 장기 수익 극대화"],
              ])
    doc.add_paragraph()

    add_para(doc, "Kelly 공식: f* = (p×b - 1) / (b - 1), Half-Kelly = f*/2")
    add_para(doc, "  p = 모델 예측 이변 확률, b = winOdds")
    add_para(doc, "  기대값이 음수이면 Kelly = 0 (배팅하지 않음) → 자동으로 나쁜 베팅 걸러냄")
    doc.add_paragraph()

    if roi is not None:
        add_para(doc, "5.3 결과 (test set 3,528건)")
        flat = roi[roi["strategy"] == "Flat"] if "strategy" in roi.columns else roi
        headers = ["Top%", "베팅 수", "적중", "적중률", "평균 배당", "ROI"]
        rows = []
        for _, r in flat.iterrows():
            rows.append([r["top_pct"], r["n_bets"], int(r["n_wins"]),
                         f'{r["hit_rate"]:.3f}', f'{r.get("avg_odds", 0):.1f}', f'{r["ROI"]:+.1f}%'])
        add_table(doc, headers, rows)
        doc.add_paragraph()

        add_para(doc, "무작위 비인기마 베팅 ROI: -26.5% (비교 기준)")
        add_para(doc, "모델 상위 20%까지는 수익, 30%부터 손실 전환.")
    doc.add_paragraph()

    # ===== 6. 신뢰구간 =====
    add_heading(doc, "6. 신뢰구간 — 이 결과를 얼마나 믿을 수 있나", level=1)
    add_para(doc, "부트스트랩 1,000회 복원추출로 ROI 분포를 만들어 95% 신뢰구간을 계산.")
    doc.add_paragraph()

    if ci is not None:
        headers = ["Top%", "Flat 중위값", "95% CI 하한", "95% CI 상한", "P(수익)", "Kelly P(수익)"]
        rows = []
        for _, r in ci.iterrows():
            rows.append([r["top_pct"], f'{r["flat_median"]:+.1f}%',
                         f'{r["flat_CI_low"]:+.1f}%', f'{r["flat_CI_high"]:+.1f}%',
                         f'{r["flat_profit_prob"]:.1f}%', f'{r["kelly_profit_prob"]:.1f}%'])
        add_table(doc, headers, rows)
        doc.add_paragraph()

    add_para(doc, "해석:", bold=True)
    add_para(doc, "  • 95% CI 하한이 전부 음수 → '100% 수익 보장'은 아님")
    add_para(doc, "  • P(수익) 80%+ (Top 5%) → 돈 벌 가능성이 손해 볼 가능성보다 4배 높음")
    add_para(doc, "  • Top 20%의 P(수익) = 87% → 가장 안정적")
    add_para(doc, "  • 표본 부족(3,528건)이 근본 원인. 3~5년 데이터로 재검증 필요")
    doc.add_paragraph()

    add_para(doc, "정직하게 말하면:", bold=True)
    add_para(doc, "  • 선별 능력은 확실하다 — 적중률 2.66% → 9.7% (상위 5%), Lift 3.6배")
    add_para(doc, "  • 수익성은 '가능성 있음'이지 '확정'이 아니다")
    add_para(doc, "  • 학술적으로는 '정보 가치(edge)가 존재하나 유의수준 5%를 달성하지 못함'")
    doc.add_paragraph()

    # ===== 7. Feature Importance =====
    add_heading(doc, "7. 어떤 변수가 중요했나", level=1)
    if fi is not None:
        top10 = fi.head(10)
        headers = ["순위", "피처", "중요도"]
        rows = []
        for i, (_, r) in enumerate(top10.iterrows(), 1):
            rows.append([str(i), r["feature"], f'{r["importance"]:.4f}'])
        add_table(doc, headers, rows)
        doc.add_paragraph()

        q_rank = fi[fi["feature"] == "q"].index[0] + 1 if "q" in fi["feature"].values else "N/A"
        add_para(doc, f"q(배당률 확률)의 순위: {q_rank}위", bold=True)
        add_para(doc, "q가 상위권 → 배당률이 이변 예측에서 여전히 핵심 정보. 나머지 피처들은 그 위에 정밀도를 더하는 역할.")
        doc.add_paragraph()
        add_para(doc, "해석: 마필의 최근 상태(직전 착순, 훈련량, 휴식일)가 핵심. 시장은 통산 성적에는 잘 반응하지만, 최근 컨디션 변화는 덜 반영 → 모델이 잡아낸 틈.")
    doc.add_paragraph()

    # ===== 8. 팀원 실험과의 비교 =====
    add_heading(doc, "8. 팀원 실험(배당률 제외)과의 비교", level=1)
    add_table(doc,
              ["항목", "팀원 실험", "이번 실험"],
              [
                  ["배당률 피처", "전부 제외", "q 포함"],
                  ["타겟", "upset_B (비인기마 입상)", "upset (비인기마 1착)"],
                  ["Best AUC", "0.6420 (RF)", "0.7420 (RF)"],
                  ["Lift@10%", "1.72", "—(적중률 기준 3.6x)"],
                  ["수익 배당", "연승(plcOdds)", "단승(winOdds)"],
                  ["수익 결론", "CI가 0 포함 → 불확실", "P(수익)=81% → 가능성 있음"],
                  ["불균형 처리", "적용 시 전부 나빠짐", "balanced 적용 (기본)"],
                  ["핵심 발견", "RF가 부스팅을 이김", "배당률+피처 결합 > 배당률 단독"],
              ])
    doc.add_paragraph()
    add_para(doc, "공통 발견: RandomForest가 가장 안정적. 신호가 약한 문제에서 부스팅보다 배깅이 나음.")
    add_para(doc, "차이점: 배당률을 넣으면 AUC가 0.64 → 0.74로 크게 올라감. 배당률은 이변 예측의 '기반'이고, 피처는 그 위에 '정밀도'를 더하는 구조.")
    doc.add_paragraph()

    # ===== 9. 한계 =====
    add_heading(doc, "9. 한계", level=1)
    add_para(doc, "1. 표본 부족 — test 3,528행 중 이변 94건. ROI 신뢰구간이 넓은 근본 원인.")
    add_para(doc, "2. 단승만 봄 — 연승(입상) 타겟으로 바꾸면 적중 수가 3~4배 늘어 CI가 좁아질 것.")
    add_para(doc, "3. 고배당 1건 의존 — 팀원 보고서와 동일한 문제. 260배급 적중 하나가 ROI를 지배.")
    add_para(doc, "4. 배당률 확정 시점 — winOdds는 경주 마감 후 확정. 실전에서는 베팅 시점과 확정 시점에 차이가 있을 수 있음.")
    add_para(doc, "5. 서울만 — 부경에서도 같은 결과가 나오는지 미확인.")
    doc.add_paragraph()

    # ===== 10. 결론 =====
    add_heading(doc, "10. 결론 및 다음 할 일", level=1)
    add_para(doc, "10.1 핵심 발견 3가지", bold=True)
    add_para(doc, "  1. 배당률+피처 결합(C)이 배당률 단독(A)보다 이변 예측이 우수하다.")
    add_para(doc, "  2. 모델 상위 5~20% 선별 베팅 시 ROI 양수 가능 (수수료 반영 후).")
    add_para(doc, "  3. 모델의 역할: '모든 경주에서 이기겠다'가 아니라 '시장이 놓치는 비인기마를 필터링하는 도구'.")
    doc.add_paragraph()

    add_para(doc, "10.2 다음 할 일", bold=True)
    add_para(doc, "  1. 연승(plcOdds) 타겟으로 재실험 → 적중 수 3~4배 증가, CI 안정화")
    add_para(doc, "  2. 인기마 붕괴 모델과 조합 → '붕괴 위험 높은 경주에서만 다크호스 베팅'")
    add_para(doc, "  3. 고배당 1건 제거 로버스트니스 체크 추가")
    add_para(doc, "  4. 피처 엔지니어링: 최근 3경주 트렌드, 기수 최근 폼, 트랙별 궁합")
    doc.add_paragraph()

    # ===== 부록 =====
    add_heading(doc, "부록 A. 실행 코드", level=1)
    code = """from sklearn.ensemble import RandomForestClassifier

# 1) 비인기마만 추출
df = df[df["pop_pct"] >= 0.5]

# 2) 피처 정의 (q 포함, 배당률 나머지 13개 제외)
C_features = [c for c in df.columns if c not in EXCLUDE_COLS]
# EXCLUDE: ID, 결과, 라벨, 배당률13개, fold

# 3) 학습
model = RandomForestClassifier(
    n_estimators=300, max_depth=10, min_samples_leaf=20,
    class_weight='balanced', random_state=42, n_jobs=-1)
model.fit(X_train, y_train)

# 4) 예측 → 순위만 사용
proba = model.predict_proba(X_test)[:, 1]
top_20pct = test.nlargest(int(len(test)*0.2), 'proba')

# 5) ROI 계산 (winOdds에 공제율 이미 반영)
roi = (top_20pct['upset'] * top_20pct['winOdds']).sum() / len(top_20pct) - 1"""

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()

    add_heading(doc, "부록 B. 실험 수치 전문", level=1)
    add_para(doc, "B.1 모델 비교 (test 3,528 / 기저율 2.66%)")
    if comp is not None:
        add_para(doc, comp.to_string(index=False))
    doc.add_paragraph()

    add_para(doc, "B.2 베팅 전략 비교")
    if roi is not None:
        add_para(doc, roi.to_string(index=False))
    doc.add_paragraph()

    add_para(doc, "B.3 신뢰구간")
    if ci is not None:
        add_para(doc, ci.to_string(index=False))
    doc.add_paragraph()

    # Save
    output_path = OUTPUT_DIR / "reports/4_이변모델/04_배당률포함_모델.docx"
    doc.save(str(output_path))
    logger.info(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
