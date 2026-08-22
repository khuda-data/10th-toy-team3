"""
13_final_report_docx.py — 이변 예측 모델 최종 전략 보고서 (docx)

12_final_strategy.py의 결과를 기반으로 상세 보고서를 생성.

실행:
    python src/upset_with_odds/13_final_report_docx.py

출력:
    results/upset_improvements/이변_예측모델_최종전략_보고서.docx
"""

import logging
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("results/upset_improvements")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load results
    strategy = pd.read_csv(OUTPUT_DIR / "final_strategy.csv") if (OUTPUT_DIR / "final_strategy.csv").exists() else None
    robustness = pd.read_csv(OUTPUT_DIR / "robustness.csv") if (OUTPUT_DIR / "robustness.csv").exists() else None
    improvements = pd.read_csv(OUTPUT_DIR / "improvement_results.csv") if (OUTPUT_DIR / "improvement_results.csv").exists() else None

    doc = Document()

    # ===== TITLE =====
    doc.add_heading("KHUDA 10기 토이프로젝트 3조", level=0)
    doc.add_heading("이변 예측 모델 — 최종 전략 보고서", level=1)
    p = doc.add_paragraph("배당률을 피처로 포함한 이변 예측 모델의 최적 베팅 전략 도출")
    p.italic = True
    doc.add_paragraph()

    # ===== 0. 요약 =====
    doc.add_heading("0. 세 줄 요약", level=1)
    doc.add_paragraph("1. winOdds 10~50배 구간에서 모델 상위 10%만 단승 베팅하면 ROI +32.0% (수수료 반영 후).")
    doc.add_paragraph("2. 부트스트랩 1,000회 중 87.9%에서 수익 발생. 무작위 베팅(-26.5%) 대비 +58.5%p 우위.")
    doc.add_paragraph("3. 모델은 '어떤 비인기마가 이길지' 완벽히 맞추는 게 아니라, '시장이 놓치는 구간의 필터링 도구'로서 가치를 가진다.")
    doc.add_paragraph()

    # ===== 1. 배경 =====
    doc.add_heading("1. 이 보고서의 위치", level=1)
    doc.add_paragraph(
        "이 보고서는 프로젝트 전체 흐름에서 마지막 단계에 해당합니다.\n\n"
        "1단계: 승률 예측 모델 설계 → 시장을 못 넘음 (AUC 0.75 vs 시장 0.81)\n"
        "2단계: 시장 대조 → 격차 0.041까지 좁혔지만 여전히 못 넘음\n"
        "3단계: 방향 전환 → 이변(비인기마 1착) 예측으로 목표 변경\n"
        "4단계: 이변 모델 선정 → RF Lift@10% 1.72배 확인\n"
        "5단계: 배당률을 피처로 포함 → AUC 0.742, ROI +25.9%\n"
        "6단계(이 보고서): 배당 구간 필터링 → 최적 전략 도출, ROI +32.0%"
    )
    doc.add_paragraph()

    # ===== 2. 실험 설계 =====
    doc.add_heading("2. 실험 설계", level=1)

    doc.add_heading("2.1 모델", level=2)
    doc.add_paragraph(
        "• 모델: RandomForest (n_estimators=300, max_depth=10, min_samples_leaf=20, class_weight='balanced')\n"
        "• 피처: q(배당률 암묵적 확률) + 말/기수/환경 피처 87개\n"
        "• 타겟: upset = (pop_pct >= 0.5) & (win == 1) — 비인기마의 단승 1착\n"
        "• 분할: 시간순 6:2:2 (train 10,425 / valid 3,439 / test 3,528)"
    )

    doc.add_heading("2.2 전처리", level=2)
    add_table(doc,
              ["단계", "처리"],
              [
                  ["서울 필터", "56,648 → 32,888행"],
                  ["비인기마 필터", "pop_pct >= 0.5 → 17,392행"],
                  ["결측치", "구조적→0, 랜덤→train 중앙값"],
                  ["다중공선성", "배당률 14개→q만 + 고상관 23개 제거"],
                  ["인코딩", "범주형 13개 → LabelEncoder"],
                  ["스케일링", "StandardScaler (train fit)"],
              ])
    doc.add_paragraph()

    doc.add_heading("2.3 핵심 질문", level=2)
    doc.add_paragraph(
        '"모델이 이변 확률이 높다고 판단한 말 중, 어떤 배당 구간에서 베팅해야 ROI가 최대화되는가?"'
    )
    doc.add_paragraph()

    # ===== 3. 6단계 개선 실험 =====
    doc.add_heading("3. 6단계 개선 실험 결과", level=1)

    doc.add_heading("3.1 Stage 1: 단승 vs 연승", level=2)
    add_table(doc,
              ["", "단승(win)", "연승(place)"],
              [
                  ["AUC", "0.7420", "0.6936"],
                  ["Top 20% ROI", "+18.0%", "-15.7%"],
                  ["P(수익)", "87.6%", "0.4%"],
              ])
    doc.add_paragraph()
    doc.add_paragraph(
        "표 읽는 법:\n"
        "• AUC는 '이변을 일으킬 말과 아닌 말을 얼마나 잘 구분하는가'입니다. 1에 가까울수록 좋고, 0.5면 무작위.\n"
        "• Top 20% ROI는 '모델이 가장 자신 있다고 한 상위 20%에만 베팅했을 때 수익률'입니다.\n"
        "• P(수익)은 '1,000번 반복 시뮬레이션 중 돈을 벌 확률'입니다.\n\n"
        "해석:\n"
        "단승은 배당이 평균 13~17배로 높아서, 적중률이 7.5%밖에 안 돼도 한 번 맞추면 돈을 많이 벌어옵니다. "
        "반면 연승은 배당이 평균 3~5배로 낮아서, 적중률이 25%나 되는데도 수수료(20%)를 이기지 못합니다.\n\n"
        "비유하면:\n"
        "• 단승 = 10번 중 1번 맞추는데, 맞추면 13배 회수 → 남는 장사\n"
        "• 연승 = 4번 중 1번 맞추는데, 맞추면 4배 회수 → 수수료 20% 떼면 적자\n\n"
        "결론: 연승은 포기하고 단승을 유지합니다."
    )
    doc.add_paragraph()

    doc.add_heading("3.2 Stage 2: 하이퍼파라미터 조정", level=2)
    doc.add_paragraph(
        "RandomForest의 설정을 바꿔봤습니다.\n\n"
        "변경 전: min_samples_leaf=20, class_weight='balanced' (소수 클래스에 가중치)\n"
        "변경 후: min_samples_leaf=50, class_weight=None (가중치 없음)\n\n"
        "결과 (연승 기준):\n"
        "• AUC: 0.6936 → 0.7101 (+0.017)\n"
        "• ROI: -15.7% → -7.4% (+8.3%p 개선)\n\n"
        "해석:\n"
        "min_samples_leaf=50은 '나무의 잎마다 최소 50개 데이터를 요구'하는 설정입니다. "
        "이렇게 하면 나무가 소수의 특이한 사례를 외우는 것(과적합)을 방지합니다. "
        "class_weight 제거는 팀원 보고서에서도 확인된 사항으로, 우리 문제에서는 "
        "가중치를 주면 순위가 바뀌지 않으면서 부작용(과적합)만 생깁니다.\n\n"
        "다만 단승에서는 이미 기존 설정으로 ROI +18%가 나왔으므로, 급하게 바꿀 필요는 없습니다."
    )
    doc.add_paragraph()

    doc.add_heading("3.3 Stage 3: 로버스트니스 (고배당 제거)", level=2)
    doc.add_paragraph(
        "로버스트니스란 '결과가 운 한두 건에 의존하는 게 아닌지' 확인하는 테스트입니다.\n\n"
        "방법: 적중한 말 중 배당이 가장 높은 것부터 순서대로 1건, 3건, 5건, 10건을 제거하고 "
        "ROI를 다시 계산합니다. 만약 1건만 빼도 수익이 손실로 뒤집히면 '운에 의존'한 것입니다.\n\n"
        "이 실험은 연승 기준으로 수행했는데, 연승은 이미 마이너스(-7.4%)라 의미가 없었습니다. "
        "단승(ROI +18%) 기준으로는 부트스트랩 1,000회로 이미 검증: P(profit) = 87.6%.\n\n"
        "즉, 단승 전략은 고배당 몇 건에 완전히 의존하는 것은 아니지만, "
        "여전히 상위 2~3건의 고배당 적중이 수익의 큰 부분을 차지합니다. "
        "이것이 '확실한 수익'이 아니라 '높은 가능성의 수익'이라고 말하는 이유입니다."
    )
    doc.add_paragraph()

    doc.add_heading("3.4 Stage 4: 배당 구간 필터 — 핵심 발견", level=2)
    doc.add_paragraph(
        "이 실험이 이번 보고서의 가장 중요한 발견입니다.\n\n"
        "'모델이 좋다고 한 말 전부에 베팅'하는 대신, "
        "'배당이 일정 수준 이상인 말에만 베팅'하면 어떻게 되는지를 테스트했습니다.\n\n"
        "연승 배당(plcOdds) 기준 결과:\n"
        "• 전체 (필터 없음): ROI -7.4% — 마이너스\n"
        "• 5배 이상만: ROI +55.5% — 갑자기 수익!\n"
        "• 10배 이상만: ROI +284.0% — 폭발적 수익\n"
        "• 20배 이상만: ROI +2934.6% — 극단적 (표본 26건으로 불안정)\n\n"
        "왜 이런 패턴이 나타나는가:\n"
        "저배당(2~5배) 말은 시장이 '어느 정도 가능성 있다'고 판단한 말입니다. "
        "이 구간에서는 시장과 모델의 판단이 비슷해서 모델이 우위를 갖기 어렵습니다.\n\n"
        "반면 고배당(10배+) 말은 시장이 '거의 안 된다'고 판단한 말입니다. "
        "하지만 모델은 '이 말은 최근 훈련을 많이 했고, 기수가 폼이 좋다'는 정보를 보고 "
        "'시장 생각보다는 가능성이 있다'고 판단합니다. 이 틈새에서 수익이 발생합니다.\n\n"
        "이 발견을 단승에 적용한 것이 4장의 최종 전략입니다."
    )
    doc.add_paragraph()

    doc.add_heading("3.5 Stage 6: 조합 전략", level=2)
    doc.add_paragraph(
        "아이디어: '인기마가 무너질 것 같은 경주'에서만 다크호스에 베팅하면 어떨까?\n\n"
        "방법:\n"
        "1. 인기마 붕괴 모델을 별도로 학습 (인기마가 기대 이하 성적을 낼 확률 예측)\n"
        "2. 경주별로 '붕괴 위험도'를 산출 (그 경주 인기마들의 붕괴 확률 중 최대값)\n"
        "3. 위험도가 높은 경주에서만 다크호스 모델의 예측을 적용\n\n"
        "결과:\n"
        "• 조합 ROI: -11.4% (단순 전략 -7.4%보다 오히려 -4%p 나쁨)\n\n"
        "왜 실패했는가:\n"
        "교재 범위(RandomForest)로는 '인기마 붕괴'를 잘 예측하지 못합니다. "
        "팀원 보고서에서도 확인된 사항으로, 인기마 붕괴의 AUC는 0.578로 거의 무작위에 가깝습니다. "
        "인기마는 모든 정보가 이미 배당에 반영되어 있어서, 남은 틈이 극히 작기 때문입니다.\n\n"
        "결론: 이 전략은 포기합니다."
    )
    doc.add_paragraph()

    # ===== 4. 최종 전략 =====
    doc.add_heading("4. 최종 전략: 배당 구간 필터링 (단승)", level=1)

    doc.add_heading("4.1 전체 결과표", level=2)
    if strategy is not None:
        headers = ["필터", "Top%", "Pool", "베팅", "적중", "적중률", "평균배당", "ROI", "CI하한", "CI상한", "P(수익)"]
        rows = []
        for _, r in strategy.iterrows():
            rows.append([
                r["filter"], r["top_pct"], r["n_pool"], r["n_bets"], r["n_wins"],
                f'{r["hit_rate"]:.3f}', f'{r["avg_odds"]:.1f}',
                f'{r["ROI"]:+.1f}%', f'{r["CI_low"]:+.1f}%', f'{r["CI_high"]:+.1f}%', f'{r["P_profit"]:.1f}%'
            ])
        add_table(doc, headers, rows)
    doc.add_paragraph()
    doc.add_paragraph(
        "표 읽는 법:\n"
        "• 필터: 어떤 배당 범위의 말만 대상으로 했는가\n"
        "• Top%: 그 범위 안에서 모델 예측 상위 몇 %만 골랐는가\n"
        "• Pool: 필터에 걸린 전체 말 수\n"
        "• 베팅: 실제로 베팅한 건수 (Pool × Top%)\n"
        "• ROI: 수익률. +면 이익, -면 손해. 수수료(20%) 이미 반영됨\n"
        "• CI하한/상한: 95% 신뢰구간. '운이 나빠도/좋아도 이 범위 안에 있다'\n"
        "• P(수익): 1,000번 시뮬레이션 중 수익이 난 비율. 높을수록 안정적\n\n"
        "핵심 패턴:\n"
        "• 필터 없이(전체) + top 20% = ROI +18%, P=87% — 이미 꽤 좋음\n"
        "• winOdds 10~50배 + top 10% = ROI +32%, P=88% — 더 좋음!\n"
        "• winOdds 30배+ 이상 = ROI 급락 — 너무 고배당은 오히려 안 좋음\n\n"
        "결론: 배당 10~50배 구간이 '스위트 스팟'입니다."
    )

    doc.add_heading("4.2 TOP 5 전략 (수익 확률 순)", level=2)
    if strategy is not None:
        top5 = strategy.sort_values("P_profit", ascending=False).head(5)
        headers = ["순위", "전략", "ROI", "P(수익)", "베팅 수", "적중"]
        rows = []
        for i, (_, r) in enumerate(top5.iterrows(), 1):
            rows.append([
                f"#{i}", f'{r["filter"]} top{r["top_pct"]}',
                f'{r["ROI"]:+.1f}%', f'{r["P_profit"]:.1f}%',
                str(int(r["n_bets"])), str(int(r["n_wins"]))
            ])
        add_table(doc, headers, rows)
    doc.add_paragraph()
    doc.add_paragraph(
        "TOP 5 해석:\n"
        "• #1(winOdds 10~50배 top 10%)이 ROI와 P(수익) 모두 가장 좋습니다.\n"
        "  237건 베팅으로 표본도 충분하고, 87.9%에서 수익이 발생합니다.\n\n"
        "• #2와 #3(필터 없음/5배+ top 20%)은 705건으로 가장 안정적이지만 ROI가 +18%로 #1보다 낮습니다.\n"
        "  '조금 덜 벌지만 더 안전한' 전략입니다.\n\n"
        "• #4(10~30배 top 5%)는 ROI +47.8%로 가장 높지만, 68건밖에 안 되서 운의 영향이 큽니다.\n"
        "  '많이 벌 수 있지만 불안정한' 전략입니다.\n\n"
        "종합하면 #1이 '수익률 + 안정성 + 표본 크기'의 균형이 가장 좋습니다."
    )

    doc.add_heading("4.3 추천 전략", level=2)
    doc.add_paragraph("winOdds 10~50배 + 모델 상위 10%", style="Intense Quote")
    add_table(doc,
              ["항목", "값"],
              [
                  ["ROI", "+32.0% (수수료 반영 후)"],
                  ["95% CI", "[-23.4%, +83.1%]"],
                  ["P(수익)", "87.9% (1,000번 중 879번 수익)"],
                  ["베팅 수", "237건"],
                  ["적중 수", "22건"],
                  ["적중률", "9.3% (전체 평균 2.66%의 3.5배)"],
                  ["평균 winOdds", "15.1배"],
                  ["무작위 ROI", "-26.5% (비교 기준)"],
                  ["모델 우위", "+58.5%p (무작위 대비)"],
              ])
    doc.add_paragraph()
    doc.add_paragraph(
        "숫자의 의미 (하나씩):\n\n"
        "• ROI +32.0%: 100만원을 투자하면 평균 132만원을 회수. 32만원 순이익.\n\n"
        "• 95% CI [-23.4%, +83.1%]: 운이 나쁘면 23% 손해, 좋으면 83% 이익. "
        "이 범위가 0을 포함하므로 '손해 가능성도 있다'는 뜻. "
        "하지만 중앙값이 +32%이므로 이익 쪽으로 크게 치우져 있음.\n\n"
        "• P(수익) 87.9%: 1,000번 반복 시뮬레이션 중 879번은 수익, 121번은 손실. "
        "즉 '돈을 벌 확률이 잃을 확률의 7배'.\n\n"
        "• 적중률 9.3%: 237번 베팅해서 22번 맞춤. 낮아 보이지만 배당이 15배라 "
        "한 번 맞추면 15만원 회수 (1만원 베팅 기준).\n\n"
        "• 무작위 ROI -26.5%: 아무나 골라서 베팅하면 100만원 → 73.5만원. "
        "모델을 쓰면 100만원 → 132만원. 모델의 가치 = +58.5%p 차이."
    )

    doc.add_heading("4.4 왜 10~50배인가", level=2)
    doc.add_paragraph(
        "• 10배 미만: 배당이 낮아서 적중해도 공제율(20%)을 이기기 어려움\n"
        "• 10~50배: 이변이 일어날 '가능성이 있으면서' 배당도 충분히 높은 스위트 스팟\n"
        "• 50배 이상: 이변 자체가 너무 희귀(적중률 ~2%)하여 모델이 구분 못 함\n"
        "• 모델의 역할: 같은 10~50배 구간 안에서 '진짜 이길 가능성이 높은 말'을 골라냄"
    )
    doc.add_paragraph()

    # ===== 5. 실전 시뮬레이션 =====
    doc.add_heading("5. 실전 시뮬레이션 (쉽게)", level=1)
    doc.add_paragraph(
        "테스트 기간: 2025년 12월 ~ 2026년 8월 (약 8개월)\n\n"
        "매 경주에서:\n"
        "  1. 비인기마(인기 하위 50%) 중 배당 10~50배인 말을 추림\n"
        "  2. 모델이 '이변 확률 높음'으로 예측한 상위 10%만 선택\n"
        "  3. 각 1만원씩 단승 베팅\n\n"
        "결과:\n"
        "  • 총 베팅: 237회 × 1만원 = 237만원\n"
        "  • 적중: 22회, 평균 배당 15.1배\n"
        "  • 총 회수: 237만원 × (1 + 0.32) ≈ 313만원\n"
        "  • 순수익: +76만원\n\n"
        "주의:\n"
        "  • 이건 과거 데이터 기반 시뮬레이션. 미래에도 동일하게 작동하는 보장 없음\n"
        "  • 1,000번 반복 시 12%에서는 손실 발생\n"
        "  • 배당은 경주 마감 시점 확정값. 실시간 변동과 차이 있을 수 있음"
    )
    doc.add_paragraph()

    # ===== 6. 왜 작동하는가 =====
    doc.add_heading("6. 왜 이 전략이 작동하는가 (해석)", level=1)
    doc.add_paragraph(
        "1. 배당률(q)이 '대략적인 인기도'를 알려준다\n"
        "   → 10~50배 구간으로 범위를 좁힘\n\n"
        "2. 나머지 피처(훈련량, 기수 컨디션, 직전 성적 등)가 '같은 배당대 안에서 차이'를 만든다\n"
        "   → 같은 15배 말이라도 최근 훈련을 많이 하고, 기수 입상률이 높은 말을 구분\n\n"
        "3. 시장(배팅자들)은 통산 성적과 인기에 잘 반응하지만,\n"
        "   '최근 몇 주간의 컨디션 변화'는 상대적으로 덜 반영한다\n"
        "   → 여기가 모델이 시장을 이기는 틈\n\n"
        "Feature Importance 상위:\n"
        "  • 직전 경주 착순 (hr_last_ord)\n"
        "  • 14일 훈련량 (train_runs_14__z)\n"
        "  • 직전 인기 백분위 (hr_last_poppct)\n"
        "  • 배당률 (q)\n"
        "  • 휴식일 (hr_rest_days__z)"
    )
    doc.add_paragraph()

    # ===== 7. 한계 =====
    doc.add_heading("7. 한계 및 주의사항", level=1)
    doc.add_paragraph(
        "1. 표본 부족\n"
        "   237건 베팅 중 22건 적중. 통계적으로 안정적이라 보기엔 적중 수가 적다.\n"
        "   95% CI가 [-23.4%, +83.1%]로 0을 포함한다 — '확실한 수익'이 아님.\n\n"
        "2. 고배당 의존\n"
        "   22건 적중 중 상위 2~3건의 고배당이 ROI를 지배한다.\n"
        "   이들을 제거하면 ROI가 크게 떨어질 수 있다.\n\n"
        "3. 과거 데이터 기반\n"
        "   시장 효율성이 매년 높아지고 있다 (2023 AUC 0.80 → 2025 AUC 0.82).\n"
        "   미래에 이 틈이 좁아질 수 있다.\n\n"
        "4. 배당 확정 시점\n"
        "   winOdds는 경주 마감 후 확정. 실전에서는 베팅 시점의 배당과 차이가 있을 수 있다.\n\n"
        "5. 서울만 검증\n"
        "   부경(부산경남)에서 동일 패턴이 재현되는지 미확인."
    )
    doc.add_paragraph()

    # ===== 8. 결론 =====
    doc.add_heading("8. 결론", level=1)
    doc.add_paragraph(
        "정직하게 답하면:\n\n"
        "• 선별 능력은 확실하다 — 적중률 2.66% → 9.3% (3.5배), Lift 3.5배\n"
        "• 수익 가능성은 높다 — P(profit) = 87.9%, 무작위 대비 +58.5%p\n"
        "• 수익이 확정은 아니다 — CI가 0을 포함, 고배당 소수에 의존\n\n"
        "학술적으로:\n"
        "  '모델에 정보 가치(edge)가 존재하며, 배당 10~50배 구간에서 가장 뚜렷하다.\n"
        "   다만 유의수준 5%를 달성하려면 3~5년치 추가 데이터가 필요하다.'\n\n"
        "발표 한 줄:\n"
        "  '시장을 정면으로 이기는 데는 실패했지만, 시장이 구조적으로 약한 구간(10~50배 비인기마)에서\n"
        "   의미 있는 정보 우위를 확인했으며, 이를 활용한 선택적 베팅이 수수료 이후에도 양의 기대값을 가진다.'"
    )
    doc.add_paragraph()

    # ===== 9. 코드 =====
    doc.add_heading("9. 실행 코드 (재현 방법)", level=1)
    code = """cd "프로젝트 경로"

# 1) 전체 파이프라인 (모델 학습 + 기본 ROI)
python src/upset_with_odds/08_full_pipeline.py

# 2) 6단계 개선 실험
python src/upset_with_odds/11_improvement_experiments.py

# 3) 최종 전략 (배당 필터 + 부트스트랩)
python src/upset_with_odds/12_final_strategy.py

# 결과 확인
start results/upset_with_odds_v2/report.html
start results/upset_improvements/final_strategy.csv"""

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()

    # ===== 부록 =====
    doc.add_heading("부록. 전략별 상세 수치", level=1)
    if strategy is not None:
        doc.add_paragraph(strategy.to_string(index=False))

    # Save
    output_path = OUTPUT_DIR / "이변_예측모델_최종전략_보고서.docx"
    doc.save(str(output_path))
    logger.info(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
